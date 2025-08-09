import pandas as pd
import os
import openpyxl

from pathlib import Path
from docx import Document
from openpyxl import Workbook

from src.resume.customizers import customize_self_introduction
from src.utils.helpers import extract_text_from_docx, extract_text_from_txt

# Replace later with a resume document creation function
def create_demo_resume_docx(filename = 'demo_resume.docx'):
    """
    Create a sample mock_resume.docx file for demonstration purposes.
    """
    new_doc = Document()
    new_doc.add_heading('Rad Navs', 0)
    new_doc.add_paragraph('Bilingual (English-Japanese) IT Manager and Senior Software Developer with 10+ years of experience in software development, team leadership, and strategic IT planning.')
    
    # Add bullet points
    p = new_doc.add_paragraph('Contact Info:')
    p.add_run('\nEmail: pradoart.jp@gmail.com').italic = True
    p.add_run('\nPhone: +81 xx-xxxx-xxxx').italic = True

    # Add a heading and skills section
    new_doc.add_heading('Skills', level=1)
    skills_list = new_doc.add_paragraph()
    skills_list.add_run('Python, PHP, Angular, AWS, Team Leadership, Strategic IT Planning').bold = True

    # Add a section for language skills
    new_doc.add_heading('Language Skills', level=1)
    languages = new_doc.add_paragraph()
    languages.add_run('English (Native), Japanese (JLPT N2)').italic = True

    # Add a work experience section
    new_doc.add_heading('Work Experience', level=1)

    # Add a table
    table = new_doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Project Name'
    hdr_cells[1].text = 'Description'

    # Save the document
    folder = "data"
    new_file_path = f"{folder}/{filename}"
    new_doc.save(new_file_path)
    print(f"Demo document created at: {new_file_path}")

# To do: Create demo resume pdf file 

# Replace later with a rirekisho document creation function
# Note: This is still NOT good enough but it is a start.
# Create a demo rirekisho document (Japanese resume)
def create_demo_rirekisho_xlsx(filename='demo_rirekisho.xlsx'):
    """
    Create a sample Japanese-style resume (rirekisho) as an Excel file
    for demonstration purposes. This uses an A3-sized layout.
    """
    try:
        # Create a new workbook and select the active sheet
        wb = Workbook()
        ws = wb.active
        ws.title = "IT Rirekisho" # type: ignore

        # --- Fill in simulated data based on a common A3 rirekisho layout ---
        # The data is spread across more columns to simulate a wider sheet
        # Personal Information (Left Side)
        ws['A1'] = "履歴書" # type: ignore
        ws['A3'] = "氏名"  # type: ignore # Name
        ws['C3'] = "紙神 ラッド" # type: ignore # Kamigami Rad (Example name)
        ws['A4'] = "フリガナ" # type: ignore # Furigana
        ws['C4'] = "カミガミ ラッド" # type: ignore # Kamigami Rad (Furigana)
        ws['A5'] = "生年月日" # type: ignore # Date of Birth
        ws['C5'] = "1990年1月1日" # type: ignore # Jan 1, 1990
        
        # Address
        ws['A7'] = "住所" # type: ignore # Address
        ws['C7'] = "東京都新宿区西新宿1-1-1" # type: ignore # Tokyo, Shinjuku
        
        # Educational History (Left Side)
        ws['A9'] = "学歴・職歴" # type: ignore # Educational & Work History
        ws['B10'] = "年" # type: ignore # Year
        ws['C10'] = "月" # type: ignore # Month
        ws['D10'] = "学歴・職歴" # type: ignore # History Details
        
        ws['B11'] = 2008 # type: ignore
        ws['C11'] = 3 # type: ignore
        ws['D11'] = "〇〇高等学校 卒業" # type: ignore # Graduated from XX High School
        
        ws['B12'] = 2012 # type: ignore
        ws['C12'] = 3 # type: ignore
        ws['D12'] = "〇〇大学 経済学部 卒業" # type: ignore # Graduated from XX University, Faculty of Economics
        
        ws['B13'] = 2012 # type: ignore
        ws['C13'] = 4 # type: ignore
        ws['D13'] = "株式会社〇〇 入社" # type: ignore # Joined XX Co., Ltd.
        
        ws['B14'] = 2018 # type: ignore
        ws['C14'] = 6 # type: ignore
        ws['D14'] = "米国法人〇〇 出向" # type: ignore # Transferred to XX USA Corporation
        
        ws['B15'] = 2023 # type: ignore
        ws['C15'] = 3 # type: ignore
        ws['D15'] = "株式会社〇〇 退職" # type: ignore # Resigned from XX Co., Ltd.
        
        ws['B16'] = 2023 # type: ignore
        ws['C16'] = 4 # type: ignore
        ws['D16'] = "株式会社〇〇 入社" # type: ignore # Joined YY Co., Ltd.

        # Licenses and Certifications (Left Side)
        ws['A18'] = "免許・資格" # type: ignore # Licenses and Certifications
        ws['B19'] = 2012 # type: ignore
        ws['C19'] = 5 # type: ignore
        ws['D19'] = "普通自動車第一種運転免許" # type: ignore # Class 1 Ordinary Vehicle Driver's License
        
        ws['B20'] = 2020 # type: ignore
        ws['C20'] = 11 # type: ignore
        ws['D20'] = "TOEIC L&R Test 900点" # type: ignore # TOEIC L&R Test Score 900

        # Self-PR (Right Side)
        # These are now in a new column to simulate the A3 width
        ws['E1'] = "自己PR" # type: ignore # Self-PR
        ws['G1'] = "私は、これまでの職務経験で培った課題解決能力とコミュニケーション能力に自信があります。特に、米国法人への出向経験では、多様な文化を持つチームメンバーと協力し、プロジェクトを成功に導きました。貴社においても、この経験を活かし、チームに貢献できると確信しております。" # type: ignore

        # Motivation for Application (Right Side)
        ws['E3'] = "志望動機" # type: ignore # Motivation for Application
        ws['G3'] = "貴社の「イノベーションを通じて社会に貢献する」という企業理念に深く共感いたしました。これまでの経験を活かし、新たな技術開発に挑戦することで、貴社の事業拡大に貢献したいと考えております。特に、〇〇分野への貴社の取り組みに強い関心を持っており、ぜひ貢献したいと考えております。" # type: ignore

        # Save the workbook
        folder = "data"
        new_file_path = f"{folder}/{filename}"
        wb.save(new_file_path)
        print(f"Demo rirekisho created at: {new_file_path}")

    except Exception as e:
        print(f"Error creating demo rirekisho: {e}")
        return False
    
def create_demo_custom_self_introduction():
    profile = """
    Name: Rad Navs
    Professional Summary: 
    Bilingual (English-Japanese) IT Manager and Senior Software Developer with 10+ 
    years of experience in software development, team leadership, and strategic IT 
    planning. Proficient in Python, PHP, Angular, and AWS, with a proven track record of 
    delivering scalable solutions for multinational teams. JLPT N2 certified with a strong 
    interest in Japanese culture and technology. I've done the whole Software Development 
    Life Cycle from Requirements definition, code development, QA testing and finally 
    delivering to clients.

    Technical Skills: 
    Programming Languages: Python, PHP, TypeScript, JavaScript, Node.js, C/C++ 
    AI / LLM: ChatGPT, Google Gemini, Perplexity
    Frameworks: Angular, Django, CodeIgniter, Bootstrap
    Cloud Services: AWS (EC2, RDS, S3, Lambda, DynamoDB), Google (GCE, GCF, BigQuery)
    Databases: MySQL, MariaDB, PostgreSQL, NoSQL
    Tools: Git, Jira, Trello, Selenium, SVN, Mantis
    Other: Agile Methodology, Scrum Master, UNIX, Shell Scripting, Embedded Systems
    """
    website = "https://www.f-keisho.co.jp/"
    job_post = """
    【京都】システムエンジニア
    仕事内容
    【職務概要】
    大手から中小企業に至るまで、食品業界の幅広いお客様への導入実績を持っている同社にて、製造業における生産管理の手法を基盤とした食品業界向け生産管理システムの自社パッケージ開発を実施いただきます。
    ※開発言語…C＃、Java、VB6など

    【職務詳細】
    お客様との直接契約による案件がほとんどです。全て自社開発となります。同社は営業職を抱えておらず、全てお客様からの紹介により案件化されます。ご提案フェーズでは、社長とエンジニア自らがプレゼンを行い、プレ段階から関与する事ができます。自ら受注したプロジェクトの要件定義段階から下流工程まで、多岐にわたる業務に携わっていただきます。食品業界というニッチな領域に特化し、裁量と自立性のある文化が根付いています。

    会社の特徴

    同社は、製造業・食品スーパーのお客様に寄り添い、現場を第一に考え、お客様の現場のニーズに誠心誠意応えています。それらの結果、豊富な経験や貴重なノウハウを得る事ができました。また、経験やノウハウを蓄えるだけでなく、お客様の更なる発展につなげるために、絶えずチャレンジを続けるプロフェッショナル集団でありたいと考えています。これからもお客様とのコミュニケーションを大切にし、課題に対して、求められている以上の提案、活動を展開しています。

    ★システム導入と並行して業務改善提案を行います★
    同社では、システム導入と並行して、根本的な業務改善にも取り組みます。システムを導入すればあらゆる課題が解決するわけではありません。システムを利用する「人」と「物」の関わり方そのものの改善を目指して、お客様と二人三脚で最適な方法を考えます。システム導入と同時に業務全体を見直しすることで、システム化のメリットを引き出します。


    募集要項

    応募条件スキル	【必須】
    ・オブジェクト指向言語（C＃、Java、VB6）の開発実務経験1年以上

    ■安定して働ける環境
    同社は有給消化率100%を誇り、残業時間も短く、エンジニアとしてスキルを伸ばせる環境というだけでなく、安心し取り組める環境です。

    下記のような方は歓迎！
    ・本業種での業務経験
    応募条件ソフト面	・セルフスターターな方
    """

    self_intro = customize_self_introduction(
        applicant_profile=profile,
        company_website=website,
        job_post=job_post
    )
    
    print("\n--- Created Custom Self-Intro Text ---")
    print(self_intro)

def create_custom_self_introduction(profile_filepath, 
                                    website,
                                    job_post_filename):
    
    profile = extract_text_from_docx(profile_filepath)
    job_post = extract_text_from_txt(job_post_filename)

    self_intro = customize_self_introduction(
        applicant_profile=profile,
        company_website=website,
        job_post=job_post
    )
    
    try:
        # Save the document
        folder = "data/custom-intro"
        fp = Path(job_post_filename)
        new_file_path = f"{folder}/{fp.name.replace('.txt', '_self_intro.docx')}"

        new_doc = Document()
        new_doc.add_paragraph(self_intro) # type: ignore

        new_doc.save(new_file_path)
        print(f"Custom Self-Intro created at: {new_file_path}")
    except Exception as e:
        print(f"An error occurred: {e}")
        return None